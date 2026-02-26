import docx
from docx.shared import Pt
import re
import glob
import os
import difflib
from inject_code import get_tasks_code

def parse_markdown_answers():
    viva_dict = {}
    commands_dict = {}
    artifact_dir = r"C:\Users\kathin\.gemini\antigravity\brain\152e1eb6-acff-40dd-a0c4-328ded5d99e6"
    
    # Parse Viva Questions
    for filename in glob.glob(os.path.join(artifact_dir, "Answers_Sheets_*.md")):
        with open(filename, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        current_q = None
        current_a = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            m = re.match(r"^\d+\.\s*\*\*(.*?)\*\*", line)
            if m:
                if current_q:
                    viva_dict[current_q.strip()] = "\n".join(current_a).strip()
                current_q = m.group(1).strip()
                current_a = []
            elif current_q and not line.startswith("## ") and not line.startswith("# "):
                current_a.append(line)
                
        if current_q:
            viva_dict[current_q.strip()] = "\n".join(current_a).strip()

    # Parse Commands
    for filename in glob.glob(os.path.join(artifact_dir, "Commands_Sheets_*.md")):
        with open(filename, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if line.startswith("|") and "**" in line:
                parts = line.split("|")
                if len(parts) >= 3:
                    op = parts[1].replace("**", "").strip()
                    cmd = parts[2].strip()
                    cmd = cmd.replace("`", "")
                    if op and op != "Operation" and op != ":---":
                        commands_dict[op] = cmd
                        
    return viva_dict, commands_dict

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def normalize_text(text):
    return re.sub(r'[^a-zA-Z0-9]', '', text).lower()

def safe_process_docx(input_path, output_path):
    viva_dict, commands_dict = parse_markdown_answers()
    tasks_code = get_tasks_code()
    doc = docx.Document(input_path)
    
    inserted_cmds = 0
    # Process commands in tables safely
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                op_text = clean_text(row.cells[0].text)
                op_norm = normalize_text(op_text)
                for dict_op, cmd_ans in commands_dict.items():
                    dict_op_norm = normalize_text(dict_op)
                    if dict_op_norm in op_norm and len(dict_op_norm) > 5:
                        if not row.cells[1].text.strip() or "Excel Formula" in row.cells[1].text or "Command" in row.cells[1].text:
                            if "Operation" not in op_text:
                                row.cells[1].text = cmd_ans
                                inserted_cmds += 1
                        elif row.cells[1].text.strip() == "":
                             row.cells[1].text = cmd_ans
                             inserted_cmds += 1

    added_code_sheets = set()
    answered_questions = set()
    inserted_ans = 0
    
    # Process paragraphs iteratively
    for i, p in enumerate(doc.paragraphs):
        p_text = clean_text(p.text)
        if not p_text:
            continue
            
        # 1. Check for Viva Questions header to inject code
        if "Viva Questions" in p_text and len(p_text) < 25:
            sheet_num = len(added_code_sheets) + 1
            sheet_key = f"Activity Sheet {sheet_num}"
            
            if sheet_key in tasks_code and sheet_key not in added_code_sheets:
                # Add code block before "Viva Questions" header
                new_p_header = docx.oxml.OxmlElement('w:p')
                p._p.addprevious(new_p_header)
                header_para = docx.text.paragraph.Paragraph(new_p_header, p._parent)
                # Instead of string newlines, use spacing properties or normal text
                header_run = header_para.add_run(f"Code: {sheet_key} Task Implementation")
                header_run.bold = True
                header_para.paragraph_format.space_before = docx.shared.Pt(12)
                header_para.paragraph_format.keep_with_next = True
                
                new_p_code = docx.oxml.OxmlElement('w:p')
                p._p.addprevious(new_p_code)
                code_para = docx.text.paragraph.Paragraph(new_p_code, p._parent)
                code_run = code_para.add_run(tasks_code[sheet_key])
                code_run.font.name = 'Consolas'
                code_para.paragraph_format.space_after = docx.shared.Pt(18)
                
                added_code_sheets.add(sheet_key)
                
        # 2. Check for Viva Question to insert Answers
        p_norm = normalize_text(p_text)
        for q, a in viva_dict.items():
            if q in answered_questions:
                continue
                
            q_norm = normalize_text(q)
            ratio = difflib.SequenceMatcher(None, q_norm, p_norm).ratio()
            
            if q_norm == p_norm or (q_norm in p_norm and len(q_norm) > 20) or (ratio > 0.90):
                new_p = docx.oxml.OxmlElement('w:p')
                p._p.addnext(new_p)
                new_para = docx.text.paragraph.Paragraph(new_p, p._parent)
                
                # Inherit paragraph indent
                if p.paragraph_format.left_indent is not None:
                    new_para.paragraph_format.left_indent = p.paragraph_format.left_indent
                
                # Clean up answer formatting (prevent random multi-newlines)
                clean_ans = a.replace("\\n", " ").strip()
                
                run_bold = new_para.add_run("Answer: ")
                run_bold.bold = True
                run_text = new_para.add_run(clean_ans)
                run_text.bold = False
                
                # Space out answers slightly using paragraph format, NOT newlines
                new_para.paragraph_format.space_before = docx.shared.Pt(6)
                new_para.paragraph_format.space_after = docx.shared.Pt(12)
                p.paragraph_format.keep_with_next = True
                
                inserted_ans += 1
                answered_questions.add(q)
                break
                
    # DO NOT globably delete empty paragraphs. Only clean specific ones if needed, 
    # but with proper spacing parameters set above, we don't even need to risk it.
    
    doc.save(output_path)
    print(f"Inserted {inserted_ans} answers, {inserted_cmds} commands, and {len(added_code_sheets)} code blocks.")

if __name__ == "__main__":
    import sys
    # Start fresh from the original document
    input_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Applied Computational Mathematics Using MATLAB (1).docx"
    output_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Answered_Final.docx"
    
    safe_process_docx(input_file, output_file)
