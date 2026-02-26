import re
import docx

def get_tasks_code():
    return {
        "Activity Sheet 1": """% MATLAB Code for Activity Sheet 1 - Array Operations
% Step 1: Definition
A = [1 2 3; 4 5 6; 7 8 9];
B = [9 8 7; 6 5 4; 3 2 1];
v = [2; 4; 6];
disp('Matrix A:'); disp(A);
disp('Matrix B:'); disp(B);
disp('Vector v:'); disp(v);

% Step 2: Addition & Subtraction
disp('A + B:'); disp(A + B);
disp('A - B:'); disp(A - B);
disp('A + 5:'); disp(A + 5); % Scalar addition

% Step 3: Multiplication Types
disp('Matrix Mult (A*B):'); disp(A*B);
disp('Element-wise Mult (A.*B):'); disp(A.*B);
disp('A * v:'); disp(A * v);

% Step 4: Transposition & Inverse
A_transpose = A';
disp('Transpose of A:'); disp(A_transpose);
C = [2 1; 5 3]; % Non-singular matrix
inv_C = inv(C);
disp('Inverse of C:'); disp(inv_C);

% Step 5: Advanced Options
disp('Dot Product (v . v):'); disp(dot(v, v));
disp('Cross Product (Vector):'); disp(cross([1 0 0], [0 1 0]));
""",
        "Activity Sheet 2": """% MATLAB Code for Activity Sheet 2 - Linear Systems
% Determine inverse types
A = [1 2; 3 4; 5 6];
% Left inverse
A_L = inv(A'*A)*A';
disp('Left Inverse (A_L):'); disp(A_L);
disp('A_L * A (Should be Identity):'); disp(A_L * A);

A2 = [1 2 3; 4 5 6];
% Right inverse
A_R = A2'*inv(A2*A2');
disp('Right Inverse (A_R):'); disp(A_R);
disp('A2 * A_R (Should be Identity):'); disp(A2 * A_R);

% Pseudo-inverse
disp('Pseudo-inverse of A:'); disp(pinv(A));

% Solve Systems
disp('Solving x = A \\ b');
A_sq = [3 2; 1 2]; b_sq = [7; 5];
disp('Square system solution:'); disp(A_sq \\ b_sq);

A_over = [1 1; 1 -1; 2 1]; b_over = [2; 0; 3];
disp('Overdetermined solution:'); disp(A_over \\ b_over);
disp('Overdetermined Residual Error:'); disp(norm(A_over*(A_over\\b_over) - b_over));

A_under = [1 2 3; 4 5 6]; b_under = [1; 2];
x_min = pinv(A_under)*b_under;
disp('Underdetermined (min norm) solution:'); disp(x_min);
""",
        "Activity Sheet 3": """% MATLAB Code for Activity Sheet 3 - Projections
% Task 1: Vector Projection
v1 = [3; 4]; v2 = [1; 0];
proj_v1_v2 = (dot(v1, v2) / dot(v2, v2)) * v2;
disp('Projection of v1 onto v2:'); disp(proj_v1_v2);

% Task 2: Subspace Projection (Least Squares)
A = [1 0; 1 1; 1 2]; b = [6; 0; 0];
% Find projection matrix P
P = A * inv(A'*A) * A';
p_hat = P * b; % Projected vector in column space
disp('Projection Matrix P:'); disp(P);
disp('Projection p_hat:'); disp(p_hat);

% Find least squares solution x
x_hat = A \\ b;
disp('Least squares parameters x_hat:'); disp(x_hat);

% Verify Error Vector Orthogonality
e = b - p_hat;
disp('Error Vector e:'); disp(e);
disp('Check A^T * e (Should be 0):'); disp(A'*e);

% Task 3: Regression Plotting
figure;
scatter([0, 1, 2], [6, 0, 0], 'o', 'filled'); hold on;
y_fit = A * x_hat;
plot([0, 1, 2], y_fit, '-r', 'LineWidth', 2);
title('Linear Regression Fit via Projection');
xlabel('x'); ylabel('b'); legend('Data Points', 'Least Squares Fit');
""",
        "Activity Sheet 4": """% MATLAB Code for Activity Sheet 4 - Eigenvalues
% Initialize Matrix
A = [4 -2; 1 1];
disp('Matrix A:'); disp(A);

% Polynomial Roots Method
p = poly(A);
roots_p = roots(p);
disp('Characteristic Polynomial Coefficients:'); disp(p);
disp('Eigenvalues via Roots:'); disp(roots_p);

% Direct Eig Method
[V, D] = eig(A);
disp('Eigenvectors (V):'); disp(V);
disp('Eigenvalues (D):'); disp(D);

% Verify AV = VD
disp('Verification AV - VD (Should be 0):');
disp(norm(A*V - V*D));

% Trace and Determinant
t = trace(A); d = det(A);
eig_sum = sum(diag(D)); eig_prod = prod(diag(D));
fprintf('Trace: %f | Sum of Eigs: %f\\n', t, eig_sum);
fprintf('Det: %f | Prod of Eigs: %f\\n', d, eig_prod);

% Diagonalization Check
fprintf('Rank of Eigenvector Matrix: %d\\n', rank(V));
if rank(V) == size(A, 1)
    disp('Matrix is Diagonalizable.');
else
    disp('Matrix is Defective.');
end
""",
        "Activity Sheet 5": """% Excel Formula Set for Activity Sheet 5
% (This activity specifically requested Excel solutions, but here is the MATLAB equivalent for verification verification of the RC circuit):
R = 10000; % 10kOhm
C = 100e-6; % 100uF
Vs = 12;
dt = 0.001; 
t = 0:dt:1;

Vc = zeros(size(t));
% Euler's Method simulation
for i = 1:length(t)-1
    dVc_dt = (Vs - Vc(i)) / (R * C);
    Vc(i+1) = Vc(i) + dVc_dt * dt;
end

% Analytical Solution
Vc_exact = Vs * (1 - exp(-t / (R*C)));

% Plotting
figure;
plot(t, Vc, 'b--', 'LineWidth', 2); hold on;
plot(t, Vc_exact, 'r-', 'LineWidth', 1);
tau = R*C;
xline(tau, 'k--', '1 Tau'); xline(3*tau, 'k--', '3 Tau');
title('RC Circuit Step Response');
xlabel('Time (s)'); ylabel('Capacitor Voltage (V)');
legend('Euler Approximation', 'Analytical Solution');
""",
        "Activity Sheet 6": """% MATLAB Code for Activity Sheet 6 - Numerical Calculus
% Time and Displacement data
t = 0:0.1:10;
% Synthetic data: s(t) = 2t^2 + sin(t)
s = 2*t.^2 + sin(t) + 0.5*randn(size(t)); % with noise

% Velocity (First derivative) using gradient
v = gradient(s, t);

% Acceleration (Second derivative)
a = gradient(v, t);

% Filtering high-frequency noise using movmean
s_smooth = movmean(s, 5);
v_smooth = gradient(s_smooth, t);
a_smooth = gradient(v_smooth, t);

% Visualization
figure;
subplot(3,1,1); plot(t, s, 'r', t, s_smooth, 'b'); title('Displacement'); legend('Noisy', 'Smoothed');
subplot(3,1,2); plot(t, v_smooth, 'k'); title('Velocity');
subplot(3,1,3); plot(t, a_smooth, 'm'); title('Acceleration');
xlabel('Time (s)');
""",
        "Activity Sheet 7": """% MATLAB Code for Activity Sheet 7 - Integration
% Define function
f = @(x) exp(-x.^2) .* sin(x);
a = 0; b = pi;

% Trapezoidal Rule (Built-in)
N = 10;
x_nodes = linspace(a, b, N);
y_nodes = f(x_nodes);
I_trap = trapz(x_nodes, y_nodes);
disp(['Trapezoidal Rule (N=10) Area: ', num2str(I_trap)]);

% Simpson's 1/3 Rule (Custom Implementation)
h = (b-a)/(N-1);
I_simp = (h/3) * (y_nodes(1) + 4*sum(y_nodes(2:2:end-1)) + 2*sum(y_nodes(3:2:end-2)) + y_nodes(end));
disp(['Simpson 1/3 Rule (N=10) Area: ', num2str(I_simp)]);

% Exact Adaptive Integration (Ground Truth)
I_exact = integral(f, a, b);
disp(['Built-in Adaptive Integral: ', num2str(I_exact)]);

% Visualization
figure;
fplot(f, [a, b], 'LineWidth', 2); hold on;
area(x_nodes, y_nodes, 'FaceAlpha', 0.2, 'EdgeColor', 'r');
title('Numerical Integration Comparison');
xlabel('x'); ylabel('f(x)');
""",
        "Activity Sheet 8": """% MATLAB Code for Activity Sheet 8 - Taylor Series
syms x y
% Define a nonlinear 1D function
f1 = sin(x) * exp(x);

% Compute explicit Taylor Polynomials iteratively
T2 = taylor(f1, x, 'ExpansionPoint', 0, 'Order', 3); % Degree 2
T4 = taylor(f1, x, 'ExpansionPoint', 0, 'Order', 5); % Degree 4
disp('4th Degree Taylor Polynomial:'); disp(T4);

% Multi-variable Function
f2 = cos(x) * sin(y);
% Multivariate Taylor Polynomial at (0, 0)
T_multi = taylor(f2, [x, y], [0, 0], 'Order', 4);
disp('Multivariate Taylor Approximation:'); disp(T_multi);

% Visualization of 1D
figure;
fplot(f1, [-2, 2], 'k', 'LineWidth', 2); hold on;
fplot(T2, [-2, 2], 'r--');
fplot(T4, [-2, 2], 'b-.');
ylim([-5, 5]);
title('Taylor Series Approximations of sin(x)e^x');
legend('Exact Function', 'Order 2', 'Order 4');
""",
        "Activity Sheet 9": """% MATLAB Code for Activity Sheet 9 - Surface Plotting
% Define domain
[X, Y] = meshgrid(-3:0.2:3, -3:0.2:3);
% Complex 3D surface equation
Z = sin(sqrt(X.^2 + Y.^2)) ./ sqrt(X.^2 + Y.^2);
Z(isnan(Z)) = 1; % Handle limit at exactly (0,0)

% Main Figure Layout
figure;

% Standard Surf
subplot(2,2,1);
surf(X, Y, Z);
title('Standard Surf Plot');
xlabel('X'); ylabel('Y'); zlabel('Z');

% Surf with lighting and shading
subplot(2,2,2);
surfl(X, Y, Z);
shading interp; colormap('copper');
title('Illuminated Surface (surfl)');

% Mesh plot
subplot(2,2,3);
mesh(X, Y, Z);
colormap('jet');
title('Wireframe Mesh Plot');

% Combined Surf and Contour (Surfc)
subplot(2,2,4);
surfc(X, Y, Z);
view(45, 30); % Adjust camera angle
title('Surface with Base Contours (surfc)');
""",
        "Activity Sheet 10": """% MATLAB Code for Activity Sheet 10 - Contour Optimization
% Objective Function: Rosenbrock's Banana Function (Scaled)
f = @(X, Y) (1 - X).^2 + 100 * (Y - X.^2).^2;
[X, Y] = meshgrid(-2:0.1:2, -1:0.1:3);
Z = f(X, Y);

figure;
% Standard contour lines
subplot(1,2,1);
contour(X, Y, Z, 30); % 30 levels
title('Standard Contour Map');
xlabel('X'); ylabel('Y');

% Filled contour with gradient vectors
subplot(1,2,2);
contourf(X, Y, Z, 20); hold on;
[FX, FY] = gradient(Z, 0.1, 0.1);
quiver(X, Y, -FX, -FY, 'k'); % Pointing down the gradient
title('Filled Contour with Steepest Descent Vectors');

% Numerical Optimization to find the minimum
obj_fun = @(x) (1 - x(1))^2 + 100 * (x(2) - x(1)^2)^2;
x0 = [-1.5, 2]; % Starting guess
[x_opt, fval] = fminsearch(obj_fun, x0);

plot(x_opt(1), x_opt(2), 'r*', 'MarkerSize', 10, 'LineWidth', 2);
disp('Optimal Minimum Found At:'); disp(x_opt);
""",
        "Activity Sheet 11": """% MATLAB Code for Activity Sheet 11 - Gradients
syms x y
f_sym = x * exp(-x^2 - y^2);

% Analytical Gradient
grad_sym = gradient(f_sym, [x, y]);
disp('Symbolic Gradient Vector:'); disp(grad_sym);

% Numerical Evaluation
[X, Y] = meshgrid(-2:0.2:2, -2:0.2:2);
Z = X .* exp(-X.^2 - Y.^2);
[FX, FY] = gradient(Z, 0.2, 0.2);

figure;
contour(X, Y, Z, 15); hold on;
% Plot normalized vector field for pure directional view
mag = sqrt(FX.^2 + FY.^2);
quiver(X, Y, FX./mag, FY./mag, 0.5, 'r');
title('Contour Map and Normalized Gradient Flow Field');
xlabel('X'); ylabel('Y');

% Demonstrate steepest ascent
startX = 0.5; startY = 1.0;
streamline(X, Y, FX, FY, startX, startY);
""",
        "Activity Sheet 12": """% MATLAB Code for Activity Sheet 12 - Hessian Dynamics
syms x y
% Define potential energy surface
V = x^2 + 2*y^2 - 0.5*x*y - 2*x;

% Calculate Symbolic Gradient and Hessian
grad_V = gradient(V, [x, y]);
H_sym = hessian(V, [x, y]);
disp('Symbolic Hessian Matrix:'); disp(H_sym);

% Find Critical Point (Set Gradient to 0)
crit_pts = solve(grad_V == [0; 0], [x, y]);
cx = double(crit_pts.x); cy = double(crit_pts.y);
fprintf('Critical Point Found at: (%f, %f)\\n', cx, cy);

% Evaluate Hessian at Critical Point
H_eval = double(subs(H_sym, [x, y], [cx, cy]));
eigs_H = eig(H_eval);
disp('Hessian Eigenvalues:'); disp(eigs_H);

% Classification
if all(eigs_H > 0)
    disp('Classification: Local Minimum (Stable)');
elseif all(eigs_H < 0)
    disp('Classification: Local Maximum (Unstable)');
else
    disp('Classification: Saddle Point');
end
""",
        "Activity Sheet 13": """% MATLAB Code for Activity Sheet 13 - Probabilities
% Normal Distribution
mu = 50; sigma = 5;
x_norm = 30:0.1:70;
pdf_norm = normpdf(x_norm, mu, sigma);

% Weibull Distribution
scale = 50; shape = 1.5;
pdf_weibull = wblpdf(x_norm, scale, shape);

figure;
plot(x_norm, pdf_norm, 'b-', 'LineWidth', 2); hold on;
plot(x_norm, pdf_weibull, 'r--', 'LineWidth', 2);
title('Probability Density Functions');
legend('Normal \mu=50, \sigma=5', 'Weibull \lambda=50, k=1.5');
xlabel('Value'); ylabel('Probability Density');

% Monte Carlo / Random Sampling
data = normrnd(mu, sigma, [1000, 1]);
figure;
histogram(data, 'Normalization', 'pdf'); hold on;
plot(x_norm, pdf_norm, 'k', 'LineWidth', 2);
title('Empirical Histogram vs Ideal Normal Distribution');
""",
        "Activity Sheet 14": """% MATLAB Code for Activity Sheet 14 - Regression
% Generate synthetic noisy data
x = linspace(1, 10, 50)';
% True relationship: y = 2.5x + 10 + noise
y = 2.5*x + 10 + 3*randn(size(x));

% Simple Linear Regression (y = X*b)
X_matrix = [ones(size(x)), x];
beta = X_matrix \\ y;
disp('Linear Regression Coefficients (Intercept, Slope):'); disp(beta');

% Fit Line
y_fit = X_matrix * beta;

% Statistics Tracking
SST = sum((y - mean(y)).^2);
SSR = sum((y_fit - mean(y)).^2);
R2 = SSR / SST;
disp(['R-Squared Value: ', num2str(R2)]);

% Visualization
figure;
scatter(x, y, 'b', 'filled'); hold on;
plot(x, y_fit, 'r-', 'LineWidth', 2);
title(['Linear Regression Fit (R^2 = ', num2str(R2), ')']);
xlabel('Input Variable X'); ylabel('Response Variable Y');
legend('Noisy Data', 'Best Fit Line');
""",
        "Activity Sheet 15": """% MATLAB Code for Activity Sheet 15 - PCA & Multivariate
% Define mean and covariance for 3 variables
mu = [10 20 5];
Sigma = [5.0  2.5  1.0;
         2.5  6.0 -1.5;
         1.0 -1.5  4.0];

% Generate Multivariate Normal Data (N=500 samples)
rng('default'); % for reproducibility
data = mvnrnd(mu, Sigma, 500);

% Perform Principal Component Analysis (PCA)
[coeff, score, latent, tsquared, explained] = pca(data);

disp('Principal Component Coefficients (Eigenvectors):');
disp(coeff);
disp('Percentage of Variance Explained by each Component:');
disp(explained);

% Visualizations
figure;
subplot(1,2,1);
scatter3(data(:,1), data(:,2), data(:,3), 10, 'b', 'filled');
title('Original 3D Data Cloud');
xlabel('V1'); ylabel('V2'); zlabel('V3');
grid on;

% Dimensionality Reduction (Plot PC1 vs PC2)
subplot(1,2,2);
scatter(score(:,1), score(:,2), 10, 'r', 'filled');
title('Projected 2D Component Space');
xlabel('Principal Component 1'); ylabel('Principal Component 2');
grid on;
"""
    }

def update_docx_with_code(input_path, output_path):
    doc = docx.Document(input_path)
    tasks_code = get_tasks_code()
    
    # Track which sheets we already added
    added_sheets = set()
    
    # We want to insert the code RIGHT BEFORE the "Viva Questions" header for each sheet
    for i, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        
        # Identify the Viva Questions header
        if "Viva Questions" in p_text and len(p_text) < 30:
            # We need to figure out which sheet this is. 
            # Easiest way is to keep track of the sheet counter based on occurrences
            sheet_num = len(added_sheets) + 1
            sheet_key = f"Activity Sheet {sheet_num}"
            
            if sheet_key in tasks_code and sheet_key not in added_sheets:
                # Add code block before this paragraph
                
                # Create a new paragraph for the header
                new_p_header = docx.oxml.OxmlElement('w:p')
                p._p.addprevious(new_p_header)
                header_para = docx.text.paragraph.Paragraph(new_p_header, p._parent)
                header_run = header_para.add_run(f"\nCode: {sheet_key} Task Implementation\n")
                header_run.bold = True
                header_para.paragraph_format.keep_with_next = True
                
                # Create a new paragraph for the actual code
                new_p_code = docx.oxml.OxmlElement('w:p')
                p._p.addprevious(new_p_code)
                code_para = docx.text.paragraph.Paragraph(new_p_code, p._parent)
                
                # Apply monospace styling to mimic a code block
                code_run = code_para.add_run(tasks_code[sheet_key] + "\n")
                code_run.font.name = 'Consolas'
                # Light grey background equivalent (shading not easily supported in simple runs, so we just use mono font)
                
                added_sheets.add(sheet_key)

    doc.save(output_path)
    print(f"Successfully injected code for {len(added_sheets)} Activity Sheets.")

# Call this specifically to fix empty paragraphs in docx
def remove_empty_paragraphs(doc):
    for p in list(doc.paragraphs):
        # Let's remove completely empty lines (like purely whitespace strings before 'Viva Questions' or in answers)
        if not p.text.strip() and not p.runs:
            p_element = p._element
            p_element.getparent().remove(p_element)

if __name__ == "__main__":
    # We will run this on the previously answered v3 document
    # and save it as v4
    import sys
    try:
        from update_docx import process_docx
    except ImportError:
        pass
        
    input_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Answered_v3.docx"
    output_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Answered_v4.docx"
    
    # Just to be safe, we open the file, clean the paragraphs, add code, and save
    doc = docx.Document(input_file)
    remove_empty_paragraphs(doc)
    doc.save(output_file)
    
    # Now inject code
    update_docx_with_code(output_file, output_file)
