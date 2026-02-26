import docx
import re
import glob
import os

def parse_markdown_answers():
    viva_dict = {}
    commands_dict = {}

    artifact_dir = r"C:\Users\kathin\.gemini\antigravity\brain\152e1eb6-acff-40dd-a0c4-328ded5d99e6"
    
    # Parse Viva Questions
    for filename in glob.glob(os.path.join(artifact_dir, "Answers_Sheets_*.md")):
        with open(filename, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        current_q = None
        current_a = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Match "1. **Question text**"
            m = re.match(r"^\d+\.\s*\*\*(.*?)\*\*", line)
            if m:
                if current_q:
                    viva_dict[current_q.strip()] = "\n".join(current_a).strip()
                current_q = m.group(1).strip()
                current_a = []
            elif current_q and not line.startswith("## ") and not line.startswith("# "):
                current_a.append(line)
                
        if current_q:
            viva_dict[current_q.strip()] = "\n".join(current_a).strip()

    # Parse Commands
    for filename in glob.glob(os.path.join(artifact_dir, "Commands_Sheets_*.md")):
        with open(filename, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            # Match "| **Operation** | `Command` |"
            if line.startswith("|") and "**" in line:
                parts = line.split("|")
                if len(parts) >= 3:
                    op = parts[1].replace("**", "").strip()
                    cmd = parts[2].strip()
                    # clean up markdown backticks for insertion
                    cmd = cmd.replace("`", "")
                    if op and op != "Operation" and op != ":---":
                        commands_dict[op] = cmd
                        
    return viva_dict, commands_dict

def clean_text(text):
    # Remove weird whitespace and normalise
    return re.sub(r'\s+', ' ', text).strip()

def normalize_text(text):
    import re
    return re.sub(r'[^a-zA-Z0-9]', '', text).lower()

def process_docx(input_path, output_path):
    viva_dict, commands_dict = parse_markdown_answers()
    doc = docx.Document(input_path)
    
    # Process tables
    inserted_cmds = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                op_text = clean_text(row.cells[0].text)
                op_norm = normalize_text(op_text)
                for dict_op, cmd_ans in commands_dict.items():
                    dict_op_norm = normalize_text(dict_op)
                    if dict_op_norm in op_norm and len(dict_op_norm) > 5:
                        # only overwrite if cell is empty or has placeholder
                        if not row.cells[1].text.strip() or "Excel Formula" in row.cells[1].text or "Command" in row.cells[1].text:
                            if "Operation" not in op_text:
                                row.cells[1].text = cmd_ans
                                inserted_cmds += 1
                        elif row.cells[1].text.strip() == "":
                             row.cells[1].text = cmd_ans
                             inserted_cmds += 1
    
    # Process paragraphs
    inserted_ans = 0
    import difflib
    answered_questions = set()
    for p in doc.paragraphs:
        p_text = clean_text(p.text)
        if not p_text:
            continue
        p_norm = normalize_text(p_text)
            
        for q, a in viva_dict.items():
            if q in answered_questions:
                continue
                
            q_norm = normalize_text(q)
            ratio = difflib.SequenceMatcher(None, q_norm, p_norm).ratio()
            
            if q_norm == p_norm or (q_norm in p_norm and len(q_norm) > 20) or (ratio > 0.90):
                # Add answer right after
                new_p = docx.oxml.OxmlElement('w:p')
                p._p.addnext(new_p)
                new_para = docx.text.paragraph.Paragraph(new_p, p._parent)
                
                # Inherit paragraph indent
                if p.paragraph_format.left_indent is not None:
                    new_para.paragraph_format.left_indent = p.paragraph_format.left_indent
                
                # Prevent page break between question and answer
                p.paragraph_format.keep_with_next = True
                
                # Bold "Answer: " only
                run_bold = new_para.add_run("Answer: ")
                run_bold.bold = True
                run_text = new_para.add_run(a)
                run_text.bold = False
                
                inserted_ans += 1
                answered_questions.add(q)
                break
                
    doc.save(output_path)
    print(f"Inserted {inserted_ans} answers and {inserted_cmds} commands.")
                
    doc.save(output_path)

if __name__ == "__main__":
    input_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Applied Computational Mathematics Using MATLAB (1).docx"
    output_file = r"c:\Users\kathin\OneDrive\Desktop\answerhelper\VAP_Answered_v3.docx"
    process_docx(input_file, output_file)
    print("Done generating VAP_Answered_v3.docx")
